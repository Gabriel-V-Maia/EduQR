import io
import math
from typing import Callable, Optional

from .models import ClassEntry, GenerationConfig, TicketTemplate
from .utils import generate_qr_bytes


def _twips(cm: float) -> int:
    return int(cm * 567)


def _set_cell_borders(cell, size: int = 12, color: str = "2D9E6B"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        t = OxmlElement(f"w:{edge}")
        t.set(qn("w:val"), "single")
        t.set(qn("w:sz"), str(size))
        t.set(qn("w:space"), "0")
        t.set(qn("w:color"), color)
        b.append(t)
    tcPr.append(b)


def _set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def _set_cell_valign(cell, val: str = "center"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def _set_cell_width(cell, cm: float):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(_twips(cm)))
    w.set(qn("w:type"), "dxa")
    tcPr.append(w)


def _set_row_height(row, cm: float):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(_twips(cm)))
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def _set_para_spacing(para, before: int = 0, after: int = 0):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement("w:spacing")
    s.set(qn("w:before"), str(before))
    s.set(qn("w:after"), str(after))
    pPr.append(s)


def _fill_ticket_cell(
    cell,
    class_name: str,
    qr_bytes: bytes,
    qr_cm: float,
    template: TicketTemplate,
):
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    GREEN = RGBColor(0x1A, 0x7A, 0x50)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)

    for p in cell.paragraphs:
        p.clear()

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p1, before=40, after=0)
    r1 = p1.add_run(template.title)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = GREEN

    if template.subtitle.strip():
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p2, before=0, after=16)
        r2 = p2.add_run(template.subtitle)
        r2.bold = False
        r2.font.size = Pt(8)
        r2.font.color.rgb = GREEN
    else:
        _set_para_spacing(p1, before=40, after=16)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p3, before=0, after=16)
    p3.add_run().add_picture(io.BytesIO(qr_bytes), width=Cm(qr_cm))

    if template.show_class_name:
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p4, before=0, after=40)
        label = f"{template.footer_prefix} {class_name}" if template.footer_prefix.strip() else class_name
        r4 = p4.add_run(label)
        r4.bold = True
        r4.font.size = Pt(10)
        r4.font.color.rgb = DARK


def generate_docx(
    classes: list[ClassEntry],
    config: GenerationConfig,
    on_progress: Optional[Callable[[int, int, str], None]] = None,
) -> None:
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    MS = "http://schemas.microsoft.com/office/word"
    sel = doc.settings.element
    compat = sel.find(".//w:compat", {"w": W})
    if compat is None:
        compat = etree.SubElement(sel, f"{{{W}}}compat")
    for cs in compat.findall(f"{{{W}}}compatSetting"):
        compat.remove(cs)
    for name, val in [
        ("compatibilityMode", "12"),
        ("overrideTableStyleFontSizeAndJustification", "1"),
        ("doNotFlipMirrorIndents", "1"),
    ]:
        cs = etree.SubElement(compat, f"{{{W}}}compatSetting")
        cs.set(f"{{{W}}}name", name)
        cs.set(f"{{{W}}}uri", MS)
        cs.set(f"{{{W}}}val", val)

    PAGE_H = 28.1
    PAGE_W = 19.4
    cols = config.cols
    rows = config.rows_per_page
    col_w = PAGE_W / cols
    row_h = PAGE_H / rows
    qr_cm = min(col_w * 0.55, row_h * 0.52)
    bpp = cols * rows
    total = len(classes)
    first_page = True

    for idx, entry in enumerate(classes):
        if on_progress:
            on_progress(idx, total, entry.display_name)

        qr_bytes = generate_qr_bytes(entry.link, config.logo_bytes)
        pages = math.ceil(entry.quantity / bpp)

        for pg in range(pages):
            if not first_page:
                doc.add_page_break()
            first_page = False

            start = pg * bpp
            end = min(start + bpp, entry.quantity)
            n = end - start
            n_rows = math.ceil(n / cols)

            table = doc.add_table(rows=n_rows, cols=cols)
            tbl = table._tbl
            tblPr = tbl.tblPr
            ts = tblPr.find(qn("w:tblStyle"))
            if ts is not None:
                tblPr.remove(ts)
            tW = OxmlElement("w:tblW")
            tW.set(qn("w:w"), str(_twips(col_w * cols)))
            tW.set(qn("w:type"), "dxa")
            tblPr.append(tW)

            bi = 0
            for ri in range(n_rows):
                row = table.rows[ri]
                _set_row_height(row, row_h)
                for ci in range(cols):
                    cell = row.cells[ci]
                    _set_cell_width(cell, col_w)
                    _set_cell_borders(cell)
                    _set_cell_margins(cell)
                    _set_cell_valign(cell, "center")
                    if bi < n:
                        _fill_ticket_cell(
                            cell,
                            entry.display_name,
                            qr_bytes,
                            qr_cm,
                            config.template,
                        )
                    else:
                        cell.paragraphs[0].clear()
                    bi += 1

    if on_progress:
        on_progress(total, total, "Salvando...")
    doc.save(config.output_path)
